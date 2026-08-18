"""Downloadable reports from a deterministic research payload.

PDF and Word are the same document the intelligence page already assembled:
orientation figures, agency and phase breakdowns, themes, firms, every award
in the evidence set, and the reading of those figures. Nothing is generated
by a model; the file is a layout of counted values.
"""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ACCENT = colors.HexColor("#12507e")
INK = colors.HexColor("#14181f")
SOFT = colors.HexColor("#5b6472")
LINE = colors.HexColor("#e3e7ed")
WARM = colors.HexColor("#b4690e")


def money(value: float | None) -> str:
    amount = float(value or 0)
    if amount >= 1e9:
        return f"${amount / 1e9:.1f}B"
    if amount >= 1e6:
        return f"${amount / 1e6:.1f}M"
    if amount >= 1e3:
        return f"${round(amount / 1e3)}K"
    return f"${round(amount)}"


def title_case(text: str) -> str:
    return re.sub(
        r"\w\S*",
        lambda m: m.group(0)[0].upper() + m.group(0)[1:].lower(),
        text or "",
    )


def slug(query: str) -> str:
    cleaned = "".join(c if c.isalnum() else "_" for c in (query or "research"))[:40]
    return cleaned.strip("_") or "research"


def lede(payload: dict) -> str:
    totals = payload.get("totals") or {}
    years = totals.get("years") or []
    span = f"between {years[0]} and {years[1]}" if len(years) == 2 else "in the available record"
    return (
        f"Across the {totals.get('awards', 0)} most closely related SBIR/STTR awards, "
        f"{money(totals.get('funding'))} in federal funding reached "
        f"{totals.get('companies', 0)} companies through {totals.get('agencies', 0)} "
        f"agencies {span}."
    )


def basis(payload: dict) -> str:
    coverage = payload.get("coverage") or {}
    text = (
        f"Evidence drawn from awards recorded between {coverage.get('first_year')} "
        f"and {coverage.get('complete_through')}."
    )
    partial = coverage.get("partial_years") or []
    if partial:
        text += (
            f" {', '.join(str(year) for year in partial)} is present in the export "
            "but incomplete, so it is excluded from totals and charts."
        )
    return text


def pdf_bytes(payload: dict) -> bytes:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.7 * inch,
        bottomMargin=0.7 * inch,
        title=f"SBIR research: {payload.get('query', '')}",
        author="SBIR Research Intelligence",
    )
    styles = _pdf_styles()
    story: list = []

    story.append(Paragraph("SBIR RESEARCH INTELLIGENCE", styles["kicker"]))
    story.append(Paragraph(title_case(payload.get("query") or "Research report"), styles["title"]))
    story.append(Paragraph(lede(payload), styles["lede"]))
    story.append(_figures_table(payload, styles))
    story.append(Paragraph(basis(payload), styles["basis"]))
    story.append(Spacer(1, 8))
    story.append(HRFlowable(width="100%", thickness=0.6, color=LINE, spaceAfter=12))

    story.append(Paragraph("Government signals", styles["section"]))
    story.extend(_distribution_block("Funding by agency", payload.get("agencies") or [], styles))
    story.extend(_distribution_block("Phase", payload.get("phases") or [], styles))
    story.extend(_distribution_block("Program", payload.get("programs") or [], styles))
    story.extend(_timeline_block(payload, styles))

    themes = payload.get("themes") or []
    story.append(Paragraph("Technology themes", styles["section"]))
    story.append(Paragraph(
        "Terms that recur across this evidence and are uncommon elsewhere in the corpus.",
        styles["hint"],
    ))
    if themes:
        story.append(Paragraph(
            " · ".join(f"<b>{_xml(t['label'])}</b> ({t['awards']})" for t in themes),
            styles["body"],
        ))
    else:
        story.append(Paragraph("No recurring uncommon terms met the threshold.", styles["hint"]))

    ecosystem = payload.get("ecosystem") or {}
    story.append(Paragraph("Companies &amp; ecosystem", styles["section"]))
    story.append(Paragraph("Recurring recipients", styles["subhead"]))
    story.append(Paragraph(
        "Firms appearing most often in this evidence, with their wider award history.",
        styles["hint"],
    ))
    story.extend(_firm_blocks(ecosystem.get("recurring") or [], styles, kind="recurring"))
    story.append(Paragraph("Phase I → Phase II progression", styles["subhead"]))
    story.append(Paragraph(
        "Projects the same firm carried from Phase I into Phase II within this technology area.",
        styles["hint"],
    ))
    story.extend(_firm_blocks(ecosystem.get("progressed") or [], styles, kind="progressed"))

    evidence = payload.get("evidence") or {}
    awards = payload.get("awards") or []
    story.append(Paragraph("Evidence", styles["section"]))
    story.append(Paragraph(
        f"The {evidence.get('size', len(awards))} awards below were selected from the "
        f"{evidence.get('considered', 0)} closest matches, after collapsing "
        f"{evidence.get('duplicates_removed', 0)} repeat filings of the same project "
        f"and limiting any one company to {evidence.get('per_company_cap', 3)}.",
        styles["hint"],
    ))
    for index, award in enumerate(awards, 1):
        story.append(_award_block(index, award, styles))

    story.append(Paragraph("Reading the evidence", styles["section_warm"]))
    story.append(Paragraph(
        "Observations follow from the figures above. They describe historical funding, "
        "which is not the same as current demand.",
        styles["hint"],
    ))
    bullets = [
        ListItem(Paragraph(_xml(sentence), styles["body"]), leftIndent=12, value="•")
        for sentence in payload.get("reading") or []
    ]
    if bullets:
        story.append(ListFlowable(bullets, bulletType="bullet", start="•", leftIndent=14, spaceBefore=2))
    story.append(Spacer(1, 14))
    story.append(Paragraph(
        "Every figure in this report is computed from SBIR.gov award records, not generated. "
        f"Assembled {datetime.now(timezone.utc).date().isoformat()}.",
        styles["basis"],
    ))

    document.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buffer.getvalue()


def docx_bytes(payload: dict) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = RGBColor(0x14, 0x18, 0x1F)

    kicker = document.add_paragraph("SBIR RESEARCH INTELLIGENCE")
    kicker.runs[0].font.size = Pt(9)
    kicker.runs[0].font.color.rgb = RGBColor(0x12, 0x50, 0x7E)
    kicker.runs[0].bold = True

    title = document.add_paragraph(title_case(payload.get("query") or "Research report"))
    title.runs[0].font.size = Pt(22)
    title.runs[0].bold = True

    document.add_paragraph(lede(payload))

    totals = payload.get("totals") or {}
    figures = document.add_table(rows=2, cols=4)
    figures.autofit = True
    labels = ("awards examined", "identified funding", "agencies", "companies")
    values = (
        str(totals.get("awards", 0)),
        money(totals.get("funding")),
        str(totals.get("agencies", 0)),
        str(totals.get("companies", 0)),
    )
    for index, (value, label) in enumerate(zip(values, labels)):
        figures.cell(0, index).text = value
        figures.cell(0, index).paragraphs[0].runs[0].bold = True
        figures.cell(0, index).paragraphs[0].runs[0].font.size = Pt(16)
        figures.cell(1, index).text = label
        figures.cell(1, index).paragraphs[0].runs[0].font.size = Pt(9)
        figures.cell(1, index).paragraphs[0].runs[0].font.color.rgb = RGBColor(0x5B, 0x64, 0x72)

    basis_p = document.add_paragraph(basis(payload))
    basis_p.runs[0].font.size = Pt(9)
    basis_p.runs[0].font.color.rgb = RGBColor(0x8B, 0x95, 0xA3)

    _docx_heading(document, "Government signals")
    _docx_distribution(document, "Funding by agency", payload.get("agencies") or [])
    _docx_distribution(document, "Phase", payload.get("phases") or [])
    _docx_distribution(document, "Program", payload.get("programs") or [])
    _docx_timeline(document, payload)

    _docx_heading(document, "Technology themes")
    document.add_paragraph(
        "Terms that recur across this evidence and are uncommon elsewhere in the corpus."
    )
    themes = payload.get("themes") or []
    if themes:
        document.add_paragraph(" · ".join(f"{t['label']} ({t['awards']})" for t in themes))
    else:
        document.add_paragraph("No recurring uncommon terms met the threshold.")

    ecosystem = payload.get("ecosystem") or {}
    _docx_heading(document, "Companies and ecosystem")
    sub = document.add_paragraph("Recurring recipients")
    sub.runs[0].bold = True
    document.add_paragraph(
        "Firms appearing most often in this evidence, with their wider award history."
    )
    _docx_firms(document, ecosystem.get("recurring") or [], kind="recurring")
    sub = document.add_paragraph("Phase I to Phase II progression")
    sub.runs[0].bold = True
    document.add_paragraph(
        "Projects the same firm carried from Phase I into Phase II within this technology area."
    )
    _docx_firms(document, ecosystem.get("progressed") or [], kind="progressed")

    evidence = payload.get("evidence") or {}
    awards = payload.get("awards") or []
    _docx_heading(document, "Evidence")
    document.add_paragraph(
        f"The {evidence.get('size', len(awards))} awards below were selected from the "
        f"{evidence.get('considered', 0)} closest matches, after collapsing "
        f"{evidence.get('duplicates_removed', 0)} repeat filings of the same project "
        f"and limiting any one company to {evidence.get('per_company_cap', 3)}."
    )
    for index, award in enumerate(awards, 1):
        _docx_award(document, index, award)

    heading = document.add_paragraph("Reading the evidence")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)
    heading.runs[0].font.color.rgb = RGBColor(0xB4, 0x69, 0x0E)
    document.add_paragraph(
        "Observations follow from the figures above. They describe historical funding, "
        "which is not the same as current demand."
    )
    for sentence in payload.get("reading") or []:
        document.add_paragraph(sentence, style="List Bullet")

    note = document.add_paragraph(
        "Every figure in this report is computed from SBIR.gov award records, not generated. "
        f"Assembled {datetime.now(timezone.utc).date().isoformat()}."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(0x8B, 0x95, 0xA3)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _set_docx_footer(document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------- PDF helpers ----------

def _pdf_styles() -> dict:
    base = getSampleStyleSheet()
    return {
        "kicker": ParagraphStyle(
            "Kicker", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=8, textColor=ACCENT, spaceAfter=6,
        ),
        "title": ParagraphStyle(
            "ReportTitle", parent=base["Title"], fontName="Helvetica-Bold",
            fontSize=18, leading=22, textColor=INK, alignment=TA_LEFT,
            spaceAfter=8,
        ),
        "lede": ParagraphStyle(
            "Lede", parent=base["Normal"], fontName="Helvetica",
            fontSize=11, leading=15, textColor=SOFT, spaceAfter=12,
        ),
        "basis": ParagraphStyle(
            "Basis", parent=base["Normal"], fontName="Helvetica",
            fontSize=8, leading=11, textColor=colors.HexColor("#8b95a3"), spaceAfter=6,
        ),
        "section": ParagraphStyle(
            "Section", parent=base["Heading2"], fontName="Helvetica-Bold",
            fontSize=11, textColor=ACCENT, spaceBefore=14, spaceAfter=6,
        ),
        "section_warm": ParagraphStyle(
            "SectionWarm", parent=base["Heading2"], fontName="Helvetica-Bold",
            fontSize=11, textColor=WARM, spaceBefore=14, spaceAfter=6,
        ),
        "subhead": ParagraphStyle(
            "Subhead", parent=base["Heading3"], fontName="Helvetica-Bold",
            fontSize=10, textColor=INK, spaceBefore=10, spaceAfter=2,
        ),
        "hint": ParagraphStyle(
            "Hint", parent=base["Normal"], fontName="Helvetica",
            fontSize=8.5, leading=11, textColor=SOFT, spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "Body", parent=base["Normal"], fontName="Helvetica",
            fontSize=9.5, leading=13, textColor=INK, alignment=TA_JUSTIFY, spaceAfter=4,
        ),
        "firm": ParagraphStyle(
            "Firm", parent=base["Normal"], fontName="Helvetica",
            fontSize=9, leading=12, textColor=INK, spaceAfter=4,
        ),
        "award_title": ParagraphStyle(
            "AwardTitle", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=10, leading=13, textColor=INK, spaceAfter=2,
        ),
        "award_meta": ParagraphStyle(
            "AwardMeta", parent=base["Normal"], fontName="Helvetica",
            fontSize=8.5, leading=11, textColor=SOFT, spaceAfter=3,
        ),
        "cell": ParagraphStyle(
            "Cell", parent=base["Normal"], fontName="Helvetica",
            fontSize=8.5, leading=11, textColor=INK,
        ),
        "cell_head": ParagraphStyle(
            "CellHead", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=8, leading=10, textColor=SOFT,
        ),
        "figure_value": ParagraphStyle(
            "FigureValue", parent=base["Normal"], fontName="Helvetica-Bold",
            fontSize=14, leading=17, textColor=INK, alignment=TA_LEFT,
        ),
        "figure_label": ParagraphStyle(
            "FigureLabel", parent=base["Normal"], fontName="Helvetica",
            fontSize=7.5, leading=10, textColor=SOFT,
        ),
    }


def _figures_table(payload: dict, styles: dict) -> Table:
    totals = payload.get("totals") or {}
    cells = [
        [Paragraph(str(totals.get("awards", 0)), styles["figure_value"]),
         Paragraph(money(totals.get("funding")), styles["figure_value"]),
         Paragraph(str(totals.get("agencies", 0)), styles["figure_value"]),
         Paragraph(str(totals.get("companies", 0)), styles["figure_value"])],
        [Paragraph("awards examined", styles["figure_label"]),
         Paragraph("identified funding", styles["figure_label"]),
         Paragraph("agencies", styles["figure_label"]),
         Paragraph("companies", styles["figure_label"])],
    ]
    table = Table(cells, colWidths=[1.7 * inch] * 4)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7f8fa")),
        ("BOX", (0, 0), (-1, -1), 0.4, LINE),
        ("INNERGRID", (0, 0), (-1, -1), 0.4, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, 0), 8),
        ("BOTTOMPADDING", (0, 1), (-1, 1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return table


def _distribution_block(title: str, rows: list[dict], styles: dict) -> list:
    heading = Paragraph(title, styles["subhead"])
    if not rows:
        return [heading, Paragraph("None recorded in this evidence.", styles["hint"])]
    data = [[
        Paragraph("Name", styles["cell_head"]),
        Paragraph("Awards", styles["cell_head"]),
        Paragraph("Funding", styles["cell_head"]),
    ]]
    for row in rows[:8]:
        data.append([
            Paragraph(_xml(row.get("name") or ""), styles["cell"]),
            Paragraph(str(row.get("awards", 0)), styles["cell"]),
            Paragraph(money(row.get("funding")), styles["cell"]),
        ])
    table = Table(data, colWidths=[4.4 * inch, 1.1 * inch, 1.3 * inch])
    table.setStyle(_plain_table_style())
    return [heading, table, Spacer(1, 8)]


def _timeline_block(payload: dict, styles: dict) -> list:
    points = [p for p in (payload.get("timeline") or {}).get("points") or [] if p.get("awards")]
    if not points:
        return [Paragraph("Activity over time", styles["subhead"]),
                Paragraph("Not enough dated awards to tabulate.", styles["hint"])]
    header = [Paragraph("Year", styles["cell_head"]),
              Paragraph("Awards", styles["cell_head"]),
              Paragraph("Funding", styles["cell_head"])]
    data = [header]
    for point in points:
        data.append([
            Paragraph(str(point["year"]), styles["cell"]),
            Paragraph(str(point["awards"]), styles["cell"]),
            Paragraph(money(point.get("funding")), styles["cell"]),
        ])
    table = Table(data, colWidths=[1.4 * inch, 1.4 * inch, 4.0 * inch])
    table.setStyle(_plain_table_style())
    return [Paragraph("Activity over time", styles["subhead"]), table, Spacer(1, 8)]


def _firm_blocks(firms: list[dict], styles: dict, kind: str) -> list:
    if not firms:
        message = (
            "No firm in this evidence has a broad award history."
            if kind == "recurring"
            else "No Phase I to Phase II progression found within this technology area."
        )
        return [Paragraph(message, styles["hint"])]
    blocks = []
    for firm in firms:
        name = title_case(firm.get("company") or "Unknown")
        if kind == "recurring":
            detail = (
                f"{firm.get('awards_here', 0)} award"
                f"{'s' if firm.get('awards_here', 0) != 1 else ''} here · "
                f"{firm.get('total_awards') or 0} overall · {money(firm.get('total_funding'))} · "
                f"{firm.get('first_year')}-{firm.get('last_year')}"
            )
        else:
            advanced = firm.get("topic_progressed") or 0
            detail = (
                f"{advanced} project{'s' if advanced != 1 else ''} advanced in this area · "
                f"{firm.get('total_awards') or 0} awards overall"
            )
        blocks.append(Paragraph(f"<b>{_xml(name)}</b><br/>{_xml(detail)}", styles["firm"]))
    return blocks


def _award_block(index: int, award: dict, styles: dict):
    tags = [award.get("agency"), award.get("branch"), award.get("phase"),
            award.get("program"), award.get("year"),
            ", ".join(part for part in (award.get("city"), award.get("state")) if part)]
    related = ""
    extra = award.get("related_awards") or 0
    if extra > 1:
        related = (
            f" · +{extra - 1} related award{'s' if extra > 2 else ''} "
            f"({money(award.get('related_funding'))} total)"
        )
    abstract = (award.get("abstract") or "").strip()
    parts = [
        Paragraph(f"{index}. {_xml(award.get('title') or 'Untitled award')}", styles["award_title"]),
        Paragraph(
            f"<b>{_xml(title_case(award.get('company') or 'Unknown'))}</b> · {money(award.get('amount'))}",
            styles["award_meta"],
        ),
        Paragraph(_xml(" · ".join(str(tag) for tag in tags if tag) + related), styles["award_meta"]),
    ]
    if abstract:
        parts.append(Paragraph(_xml(abstract), styles["body"]))
    parts.append(Spacer(1, 6))
    return KeepTogether(parts)


def _plain_table_style() -> TableStyle:
    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f7f8fa")),
        ("LINEBELOW", (0, 0), (-1, 0), 0.6, LINE),
        ("LINEBELOW", (0, 1), (-1, -1), 0.3, LINE),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ])


def _footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setStrokeColor(LINE)
    canvas.setLineWidth(0.4)
    canvas.line(0.75 * inch, 0.5 * inch, letter[0] - 0.75 * inch, 0.5 * inch)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(SOFT)
    canvas.drawString(0.75 * inch, 0.32 * inch, "SBIR.gov award records  ·  Historical funding is not current demand")
    canvas.drawRightString(letter[0] - 0.75 * inch, 0.32 * inch, str(doc.page))
    canvas.restoreState()


def _xml(text) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ---------- Word helpers ----------

NAVY = RGBColor(0x12, 0x50, 0x7E)
MUTED = RGBColor(0x5B, 0x64, 0x72)


def _docx_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.runs[0].bold = True
    paragraph.runs[0].font.size = Pt(14)
    paragraph.runs[0].font.color.rgb = NAVY


def _docx_distribution(document: Document, title: str, rows: list[dict]) -> None:
    heading = document.add_paragraph(title)
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(11)
    if not rows:
        document.add_paragraph("None recorded in this evidence.")
        return
    table = document.add_table(rows=1 + min(len(rows), 8), cols=3)
    table.style = "Table Grid"
    headers = ("", "Awards", "Funding")
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for offset, row in enumerate(rows[:8], 1):
        table.rows[offset].cells[0].text = str(row.get("name") or "")
        table.rows[offset].cells[1].text = str(row.get("awards", 0))
        table.rows[offset].cells[2].text = money(row.get("funding"))
    document.add_paragraph("")


def _docx_timeline(document: Document, payload: dict) -> None:
    heading = document.add_paragraph("Activity over time")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(11)
    points = [p for p in (payload.get("timeline") or {}).get("points") or [] if p.get("awards")]
    if not points:
        document.add_paragraph("Not enough dated awards to tabulate.")
        return
    table = document.add_table(rows=1 + len(points), cols=3)
    table.style = "Table Grid"
    for index, label in enumerate(("Year", "Awards", "Funding")):
        table.rows[0].cells[index].text = label
        table.rows[0].cells[index].paragraphs[0].runs[0].bold = True
    for offset, point in enumerate(points, 1):
        table.rows[offset].cells[0].text = str(point["year"])
        table.rows[offset].cells[1].text = str(point["awards"])
        table.rows[offset].cells[2].text = money(point.get("funding"))
    document.add_paragraph("")


def _docx_firms(document: Document, firms: list[dict], kind: str) -> None:
    if not firms:
        message = (
            "No firm in this evidence has a broad award history."
            if kind == "recurring"
            else "No Phase I to Phase II progression found within this technology area."
        )
        document.add_paragraph(message)
        return
    for firm in firms:
        name = document.add_paragraph(title_case(firm.get("company") or "Unknown"))
        name.runs[0].bold = True
        if kind == "recurring":
            detail = (
                f"{firm.get('awards_here', 0)} award"
                f"{'s' if firm.get('awards_here', 0) != 1 else ''} here · "
                f"{firm.get('total_awards') or 0} overall · {money(firm.get('total_funding'))} · "
                f"{firm.get('first_year')}-{firm.get('last_year')}"
            )
        else:
            advanced = firm.get("topic_progressed") or 0
            detail = (
                f"{advanced} project{'s' if advanced != 1 else ''} advanced in this area · "
                f"{firm.get('total_awards') or 0} awards overall"
            )
        paragraph = document.add_paragraph(detail)
        paragraph.runs[0].font.size = Pt(10)
        paragraph.runs[0].font.color.rgb = MUTED


def _docx_award(document: Document, index: int, award: dict) -> None:
    title = document.add_paragraph(f"{index}. {award.get('title') or 'Untitled award'}")
    title.runs[0].bold = True
    company = title_case(award.get("company") or "Unknown")
    document.add_paragraph(f"{company}  ·  {money(award.get('amount'))}")
    tags = [award.get("agency"), award.get("branch"), award.get("phase"),
            award.get("program"), award.get("year"),
            ", ".join(part for part in (award.get("city"), award.get("state")) if part)]
    meta = document.add_paragraph(" · ".join(str(tag) for tag in tags if tag))
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = MUTED
    extra = award.get("related_awards") or 0
    if extra > 1:
        related = document.add_paragraph(
            f"+{extra - 1} related award{'s' if extra > 2 else ''} · "
            f"{money(award.get('related_funding'))} total"
        )
        related.runs[0].font.size = Pt(9)
        related.runs[0].font.color.rgb = RGBColor(0xB4, 0x69, 0x0E)
    abstract = (award.get("abstract") or "").strip()
    if abstract:
        document.add_paragraph(abstract)


def _set_docx_footer(document: Document) -> None:
    footer = document.sections[0].footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.text = "SBIR.gov award records  ·  Historical funding is not current demand"
    paragraph.runs[0].font.size = Pt(8)
    paragraph.runs[0].font.color.rgb = MUTED
    # PAGE field so Word numbers pages itself.
    run = paragraph.add_run("    ")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "end")
    run._r.append(fld)
